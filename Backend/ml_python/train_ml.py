#!/usr/bin/env python3
"""
Train the SriLankanFoodMLSystem end-to-end
Usage:
    python train_ml.py <categorization_docx> <recipes_docx>
"""
import os
import sys
import json
import re

import pandas as pd
import numpy as np
import tensorflow as tf
from tensorflow.keras.models import Model
from tensorflow.keras.layers import Dense, Dropout, Input
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.cluster import KMeans
from sklearn.ensemble import RandomForestClassifier
import joblib
from docx import Document

class SriLankanFoodMLSystem:
    def __init__(self):
        self.models_dir = 'ml_models'
        self.data_dir   = 'ml_data'
        self._make_dirs()

    def _make_dirs(self):
        os.makedirs(self.models_dir, exist_ok=True)
        os.makedirs(self.data_dir, exist_ok=True)

    def extract_data_from_docx(self, cat_file, rec_file):
        print(f"\n📂 Loading categorization file: {cat_file}")
        doc_cat = Document(cat_file)
        print("  → Total paragraphs in categorization doc:", len(doc_cat.paragraphs))
        cats = self._extract_categories(cat_file)
        print("🔍 Categories found:", {k: len(v) for k, v in cats.items()})

        print(f"\n📂 Loading recipes file: {rec_file}")
        doc_rec = Document(rec_file)
        print("  → Total paragraphs in recipes doc:", len(doc_rec.paragraphs))
        recs = self._extract_recipes(rec_file)
        print("🔍 Recipes found:", len(recs), "– sample keys:", list(recs.keys())[:5])

        df = self._build_dataset(cats, recs)
        print("🔍 Built dataset rows:", len(df))
        out_csv = os.path.join(self.data_dir, 'sri_lankan_foods.csv')
        df.to_csv(out_csv, index=False)
        print(f"✅ Saved dataset to {out_csv}")
        return df

    def _extract_categories(self, path):
        doc = Document(path)
        print("\n   [DEBUG] _extract_categories reading paragraphs:")
        cats = {'high_blood_pressure': [], 'diabetes': [], 'both': []}
        cur = None

        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            print(f"     ¶{i:02d}:", repr(t))

            # break on blank line
            if not t:
                cur = None
                continue

            # header detection
            if re.match(r'^Best\s+for\s+High\s+Blood\s+Pressure\b', t, re.IGNORECASE):
                cur = 'high_blood_pressure'
                print("       → Set current = high_blood_pressure")
                continue
            if re.match(r'^Best\s+for\s+Diabetes\b', t, re.IGNORECASE):
                cur = 'diabetes'
                print("       → Set current = diabetes")
                continue
            if re.match(r'^Best\s+for\s+Both\b', t, re.IGNORECASE):
                cur = 'both'
                print("       → Set current = both")
                continue

            # skip descriptive lines
            if 'recipes' in t.lower():
                continue

            # under a category, any non-blank paragraph is an item
            if cur:
                cats[cur].append(t)
                print(f"       • Added to {cur!r}:", t)

        print("   [DEBUG] Final categories dict:", cats)
        return cats

    def _extract_recipes(self, path):
        paras = [p.text.strip() for p in Document(path).paragraphs if p.text.strip()]
        print("\n   [DEBUG] _extract_recipes got", len(paras), "non-empty paragraphs")
        recipes = {}
        current = None
        section = None

        for idx, line in enumerate(paras):
            print(f"     L{idx:02d}:", repr(line))
            m = re.match(r'^(\d+)\.\s*(.+)', line)
            if m:
                if current:
                    recipes[current['name']] = current
                    print(f"       → Saved recipe: {current['name']}")
                current = {
                    'name': m.group(2).strip(),
                    'ingredients': [],
                    'instructions': [],
                    'portion_size': ''
                }
                print("       → New recipe start:", current['name'])
                section = None
                continue

            if line.lower().startswith('ingredients'):
                section = 'ingredients'
                print("       → Entering INGREDIENTS section")
                continue
            if line.lower().startswith('preparation'):
                section = 'preparation'
                print("       → Entering PREPARATION section")
                continue
            if line.lower().startswith('portion size'):
                pm = re.search(r'Portion Size[:\s]*(.+)', line, re.IGNORECASE)
                if pm:
                    current['portion_size'] = pm.group(1).strip()
                    print("       → Portion size:", current['portion_size'])
                section = None
                continue

            if section == 'ingredients' and line.startswith('•'):
                ingr = line.lstrip('•\t ').strip()
                current['ingredients'].append(ingr)
                print("         • Ingredient:", ingr)
                continue

            if section == 'preparation' and re.match(r'^\d+\.', line):
                instr = re.sub(r'^\d+\.\s*', '', line).strip()
                current['instructions'].append(instr)
                print("         • Instruction:", instr)
                continue

        if current:
            recipes[current['name']] = current
            print(f"       → Saved last recipe: {current['name']}")
        print("   [DEBUG] Recipe keys:", list(recipes.keys()))
        return recipes

    def _build_dataset(self, cats, recs):
        nut_db = self._nutrition_db()
        rows, fid = [], 1
        for cat, names in cats.items():
            for nm in names:
                rc = recs.get(nm, {
                    'name': nm,
                    'ingredients': [f"Traditional {nm} method"],
                    'instructions': ["Follow traditional method"],
                    'portion_size': "1 serving"
                })
                nut = self._estimate_nutrition(rc, nm, nut_db)
                row = {
                    'id': fid,
                    'name': nm,
                    'health_category': cat,
                    'ingredients': json.dumps(rc['ingredients']),
                    'instructions': json.dumps(rc['instructions']),
                    'portion_size': rc['portion_size'],
                    **nut
                }
                rows.append(row)
                fid += 1
        return pd.DataFrame(rows)

    def _nutrition_db(self):
        return {
            'kurakkan': {'protein':7.3,'fat':1.3,'carbs':72.6,'potassium':408,'fiber':11.5,'sodium':0.1,'magnesium':2.8},
            'coconut':  {'protein':3.3,'fat':33.5,'carbs':15.2,'potassium':356,'fiber':9.0,'sodium':0.2,'magnesium':3.2},
            'rice':     {'protein':2.7,'fat':0.3,'carbs':28.0,'potassium':115,'fiber':0.4,'sodium':0.1,'magnesium':0.6},
        }

    def _estimate_nutrition(self, recipe, name, db):
        base = {'protein':2.0,'fat':1.0,'carbs':15.0,'potassium':300,'fiber':2.5,'sodium':0.5,'magnesium':0.8}
        txt  = ' '.join(recipe['ingredients']).lower()
        nm   = name.lower()
        for ingr, vals in db.items():
            if ingr in txt or ingr in nm:
                w = 0.3
                for k in base:
                    base[k] += vals.get(k, 0) * w
        base['calories']       = int(base['protein']*4 + base['fat']*9 + base['carbs']*4)
        base['cooking_time']   = 20
        base['health_score']   = min(100, 60 + (base['fiber']>4)*15 + (base['potassium']>400)*15)
        base['glycemic_index'] = 50
        return base

    def train_all_models(self, df):
        print("🤖 Training models…")
        feats = [
            'protein','fat','carbs','potassium','fiber','sodium','magnesium',
            'calories','cooking_time','glycemic_index'
        ]
        if df.empty:
            raise ValueError("Dataset is empty—check your DOCX parsing.")
        X = df[feats].values

        # Content-based
        scaler_cb = StandardScaler().fit(X)
        X_cb      = scaler_cb.transform(X)
        y_bp      = (df['health_category'].isin(['high_blood_pressure','both'])).astype(int)
        y_db      = (df['health_category'].isin(['diabetes','both'])).astype(int)
        y_hs      = df['health_score'] / 100.0

        inp = Input(shape=(X.shape[1],))
        h1  = Dense(64, activation='relu')(inp)
        d1  = Dropout(0.3)(h1)
        h2  = Dense(32, activation='relu')(d1)
        d2  = Dropout(0.2)(h2)
        bp_out = Dense(1, activation='sigmoid', name='blood_pressure')(d2)
        db_out = Dense(1, activation='sigmoid', name='diabetes')(d2)
        hs_out = Dense(1, activation='sigmoid', name='health_score')(d2)

        model_cb = Model(inp, [bp_out, db_out, hs_out])
        model_cb.compile(
            optimizer='adam',
            loss={
                'blood_pressure': 'binary_crossentropy',
                'diabetes':       'binary_crossentropy',
                'health_score':   'mse'
            },
            metrics={
                'blood_pressure': ['accuracy'],
                'diabetes':       ['accuracy'],
                'health_score':   ['mae']
            }
        )

        model_cb.fit(
            X_cb, [y_bp, y_db, y_hs],
            epochs=30, batch_size=8, validation_split=0.2, verbose=0
        )
        model_cb.save(os.path.join(self.models_dir, 'content_based_model.h5'))
        joblib.dump(scaler_cb, os.path.join(self.models_dir, 'content_scaler.pkl'))
        print(" • Content-based model trained")

        # Clustering
        scaler_cl = StandardScaler().fit(X)
        X_cl      = scaler_cl.transform(X)
        k         = min(8, len(df)//3)
        kmeans    = KMeans(n_clusters=k, random_state=42, n_init=10).fit(X_cl)
        joblib.dump(kmeans, os.path.join(self.models_dir, 'kmeans.pkl'))
        joblib.dump(scaler_cl, os.path.join(self.models_dir, 'kmeans_scaler.pkl'))
        print(" • Clustering model trained")

        # Health classifier
        le = LabelEncoder().fit(df['health_category'])
        y  = le.transform(df['health_category'])
        rf = RandomForestClassifier(n_estimators=50, random_state=42).fit(X, y)
        joblib.dump(rf, os.path.join(self.models_dir, 'rf_health.pkl'))
        joblib.dump(le, os.path.join(self.models_dir, 'health_le.pkl'))
        print(" • Health classifier trained")

        # Metadata
        meta = {
            'total_rows': len(df),
            'categories': df['health_category'].value_counts().to_dict(),
            'features': feats
        }
        with open(os.path.join(self.models_dir, 'metadata.json'), 'w') as f:
            json.dump(meta, f, indent=2)
        print("✅ All models trained!")

def main():
    if len(sys.argv) != 3:
        print("Usage: python train_ml.py <categorization.docx> <recipes.docx>")
        sys.exit(1)

    cat_file = sys.argv[1]
    rec_file = sys.argv[2]

    ml_system = SriLankanFoodMLSystem()
    df        = ml_system.extract_data_from_docx(cat_file, rec_file)
    ml_system.train_all_models(df)

if __name__ == "__main__":
    main()
